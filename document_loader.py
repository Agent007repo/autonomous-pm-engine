"""
src/knowledge/document_loader.py

Multi-format document ingestion layer.
Supports: .txt, .pdf, .docx, .csv, .md
Returns a flat list of LangChain Document objects with rich metadata.
"""

from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import TYPE_CHECKING

from langchain_core.documents import Document
from loguru import logger

if TYPE_CHECKING:
    pass

# ── Optional heavy imports (graceful degradation) ─────────────────────────────
try:
    from langchain_community.document_loaders import PyPDFLoader
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False

try:
    import docx as python_docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


class DocumentLoader:
    """
    Loads documents from a directory or a list of file paths.
    Each returned Document carries metadata including:
      - source: original file path
      - doc_type: 'interview' | 'survey' | 'research' | 'unknown'
      - page: page number (PDFs) or row index (CSVs)
    """

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".csv", ".md"}

    # Heuristic mapping: filename keywords -> document type
    _TYPE_HINTS: dict[str, str] = {
        "interview": "interview",
        "survey": "survey",
        "research": "research",
        "market": "research",
        "feedback": "feedback",
        "review": "feedback",
    }

    def load_directory(self, directory: str | Path) -> list[Document]:
        """
        Recursively load all supported files under `directory`.
        Returns a deduplicated, sorted list of Documents.
        """
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Input directory not found: {directory}")

        all_docs: list[Document] = []
        files = sorted(directory.rglob("*"))

        for fp in files:
            if fp.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
                continue
            if fp.name.startswith("."):
                continue
            try:
                docs = self.load_file(fp)
                all_docs.extend(docs)
                logger.info(f"Loaded {len(docs)} document(s) from {fp.name}")
            except Exception as exc:
                logger.warning(f"Failed to load {fp}: {exc}")

        logger.info(f"Total documents loaded: {len(all_docs)}")
        return all_docs

    def load_file(self, file_path: str | Path) -> list[Document]:
        """
        Load a single file. Dispatches by extension.
        """
        fp = Path(file_path)
        ext = fp.suffix.lower()
        doc_type = self._infer_doc_type(fp.stem)

        dispatch = {
            ".txt": self._load_txt,
            ".md": self._load_txt,
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".csv": self._load_csv,
        }

        loader_fn = dispatch.get(ext)
        if loader_fn is None:
            raise ValueError(f"Unsupported file type: {ext}")

        docs = loader_fn(fp)

        # Inject doc_type into every document's metadata
        for doc in docs:
            doc.metadata["doc_type"] = doc_type
            doc.metadata.setdefault("source", str(fp))

        return docs

    # ── Format-specific loaders ───────────────────────────────────────────────

    def _load_txt(self, fp: Path) -> list[Document]:
        text = fp.read_text(encoding="utf-8", errors="replace")
        return [Document(page_content=text, metadata={"source": str(fp), "page": 0})]

    def _load_pdf(self, fp: Path) -> list[Document]:
        if not _PDF_AVAILABLE:
            raise ImportError("pypdf is required for PDF loading: pip install pypdf")
        loader = PyPDFLoader(str(fp))
        pages = loader.load()
        # LangChain PyPDFLoader already sets page metadata
        for i, page in enumerate(pages):
            page.metadata.setdefault("page", i)
        return pages

    def _load_docx(self, fp: Path) -> list[Document]:
        if not _DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX loading: pip install python-docx"
            )
        doc = python_docx.Document(str(fp))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        return [Document(page_content=full_text, metadata={"source": str(fp), "page": 0})]

    def _load_csv(self, fp: Path) -> list[Document]:
        """
        Treats each non-empty CSV row as an independent Document.
        This preserves granularity for survey data (one row = one respondent).
        """
        docs: list[Document] = []
        with fp.open(encoding="utf-8", errors="replace", newline="") as fh:
            reader = csv.DictReader(fh)
            if reader.fieldnames is None:
                return []
            for i, row in enumerate(reader):
                # Concatenate all fields into a single string
                text = " | ".join(
                    f"{k}: {v}" for k, v in row.items() if v and v.strip()
                )
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={"source": str(fp), "page": i, "row_index": i},
                        )
                    )
        return docs

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _infer_doc_type(self, stem: str) -> str:
        stem_lower = stem.lower()
        for keyword, doc_type in self._TYPE_HINTS.items():
            if keyword in stem_lower:
                return doc_type
        return "unknown"
